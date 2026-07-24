from fastapi import FastAPI, HTTPException, status, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from groq import Groq
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
import os
import json
import re
import io
import base64
import httpx

# ── env + client ────────────────────────────────────────────────────────────
load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")

# ── app + rate limiter ───────────────────────────────────────────────────────
limiter = Limiter(key_func=get_remote_address)
app = FastAPI()
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── OCR config (cloud vision, no local model download) ────────────────────────
# EasyOCR was removed: it downloads ~100MB of weights into $HOME/.EasyOCR at
# runtime, which fails on Render's ephemeral/read-only-ish filesystem
# ("No such file or directory: /opt/render/.EasyOCR//model/temp.zip") and would
# OOM the 512MB free instance anyway once torch loads.
VISION_MODELS = [
    os.getenv("GROQ_VISION_MODEL", "qwen/qwen3.6-27b"),
    "meta-llama/llama-4-scout-17b-16e-instruct",
    "meta-llama/llama-4-maverick-17b-128e-instruct",
]

OCR_PROMPT = (
    "Transcribe ALL text visible in this image, exactly as written. "
    "Include sender names, phone numbers, links, amounts, dates and button labels. "
    "Do not summarise, translate, explain or add any commentary. "
    "Output only the transcribed text. "
    "If the image contains no readable text, output exactly: NO_TEXT_FOUND"
)

MAX_UPLOAD_BYTES = 8 * 1024 * 1024   # 8MB hard cap on uploads
MAX_IMAGE_SIDE = 1600                # downscale before sending to the model

# ── request models ────────────────────────────────────────────────────────────
class MessageRequest(BaseModel):
    message: str

class ReportRequest(BaseModel):
    message_preview: str
    scam_type: str | None
    risk_score: int
    verdict: str
    red_flags: str

class ChatRequest(BaseModel):
    session_id: str
    message: str

# ── prompts ──────────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """You are a scam detection expert specializing in fake job messages circulated on WhatsApp and Telegram in India.

SCORING GUIDE — be strict and calibrated:
- 0-20: Clearly legitimate. Professional tone, known company, no payment request, verifiable contact.
- 21-40: Minor concerns only. Slightly vague but no active red flags.
- 41-65: Suspicious. Multiple vague claims, unverifiable, but no direct scam pattern confirmed.
- 66-85: Likely scam. Clear red flags: urgency, guaranteed income, registration fee mention, no company name.
- 86-100: Definite scam. Explicit payment demand, impersonation, guaranteed job with no interview.

SCAM PATTERNS to detect:
- Registration Fee Scam: asks for upfront payment to register or get the job
- Guaranteed Job Scam: promises job without interview or qualification check
- Work From Home Scam: vague WFH role, unrealistic daily/monthly income (e.g. "50k/month ghar baithe")
- Data Harvesting Scam: asks for Aadhaar, PAN, bank details early in the process
- Impersonation Scam: claims to be from TCS, Infosys, govt, etc. but contact is WhatsApp/personal number
- Other: doesn't fit above but still suspicious

SAFE SIGNALS that reduce score:
- Official company email domain (not Gmail/Yahoo)
- Specific job title and department
- Interview process mentioned
- No payment requested
- Verifiable company name with website

Respond ONLY with a valid JSON object. No explanation, no markdown, no backticks."""

FEW_SHOT_EXAMPLES = """
Example 1 — Definite scam:
Message: "Urgent hiring! Work from home. Earn Rs 800 per hour. No experience needed. Registration fee Rs 500 only. Call now: 9876543210"
{"risk_score": 94, "verdict": "Likely Scam", "scam_type": "Registration Fee Scam", "red_flags": ["Upfront registration fee demanded", "Unrealistic hourly income claim", "No company name or website", "Urgency pressure tactic", "Personal phone number only"], "safe_signals": [], "advice": "Never pay any fee to get a job — legitimate employers do not charge candidates for registration or training."}

Example 2 — Legitimate:
Message: "Dear candidate, this is to inform you that your application for the role of Software Engineer at Infosys BPM has been shortlisted. Please attend the interview on 28th May at our Bangalore office, Electronic City Phase 1. Carry your resume and govt ID. No charges apply. HR team: hr.infosys@infosysbpm.com"
{"risk_score": 12, "verdict": "Safe", "scam_type": null, "red_flags": [], "safe_signals": ["Official company email domain used", "Specific office location provided", "No payment requested", "Named company with verifiable identity", "Interview process mentioned"], "advice": "This message appears legitimate. Verify by calling Infosys BPM directly using their official website number before attending."}

Example 3 — Suspicious:
Message: "Hi, we found your resume on Naukri. We have openings for data entry work from home. Salary 25,000/month. Interested? Reply YES."
{"risk_score": 48, "verdict": "Suspicious", "scam_type": "Work From Home Scam", "red_flags": ["No company name provided", "Vague job description", "Unsolicited contact claiming to have your resume"], "safe_signals": ["No payment mentioned", "Salary is realistic"], "advice": "Ask for the company name, official website, and HR email before proceeding — do not share personal documents yet."}
"""

CHAT_SYSTEM_PROMPT = """You are ScamShield Assistant — an expert on job scams, fraud detection, and employment safety in India.

You ONLY answer questions related to:
- Job scams and fraud detection
- How to verify if a job offer is legitimate
- Specific scam patterns (registration fee, impersonation, WFH scams, etc.)
- What to do if someone has been scammed
- How to report scams in India
- General job search safety tips
- Explaining risk scores or analysis results from ScamShield

If someone asks about ANYTHING else (coding, general knowledge, entertainment, math, etc.), respond with:
"I'm ScamShield Assistant — I can only help with job scam detection and employment fraud questions. Please ask me something related to that topic."

Always be helpful, clear, and provide actionable advice. Use simple language that a non-technical person can understand.
When relevant, mention cybercrime.gov.in and helpline 1930 for reporting scams in India."""

# ── demo fallback cache ───────────────────────────────────────────────────────
DEMO_CACHE = {
    "scam": {
        "risk_score": 94, "verdict": "Likely Scam", "scam_type": "Registration Fee Scam",
        "red_flags": ["Upfront registration fee demanded", "Unrealistic income promise", "No company name or website", "Urgency pressure tactic", "Personal WhatsApp number only"],
        "safe_signals": [], "advice": "Never pay any fee to get a job — legitimate employers do not charge candidates.",
    },
    "suspicious": {
        "risk_score": 52, "verdict": "Suspicious", "scam_type": "Work From Home Scam",
        "red_flags": ["No company name provided", "Vague job description", "Unsolicited contact"],
        "safe_signals": ["No payment mentioned", "Salary is realistic"],
        "advice": "Ask for the company name and official website before sharing any personal details.",
    },
    "safe": {
        "risk_score": 12, "verdict": "Safe", "scam_type": None,
        "red_flags": [],
        "safe_signals": ["Official company email domain used", "Specific office location provided", "No payment requested", "Interview process mentioned"],
        "advice": "This message appears legitimate. Verify by calling the company directly using their official website number.",
    },
}

SCAM_KEYWORDS = ["fee", "register", "registration", "pay", "payment", "deposit", "guaranteed", "ghar baithe", "घर बैठे", "premium", "urgent", "limited seats", "apply now", "call now", "whatsapp now", "no interview", "no experience", "earn daily", "part time earn"]
SAFE_KEYWORDS = ["interview", "office", "hr@", ".com email", "bring resume", "shortlisted", "scheduled", "department", "joining date"]

def fallback_response(message: str) -> dict:
    msg = message.lower()
    scam_hits = sum(1 for k in SCAM_KEYWORDS if k in msg)
    safe_hits = sum(1 for k in SAFE_KEYWORDS if k in msg)
    if scam_hits >= 2: return DEMO_CACHE["scam"]
    elif safe_hits >= 2: return DEMO_CACHE["safe"]
    else: return DEMO_CACHE["suspicious"]

# ── OCR helpers ───────────────────────────────────────────────────────────────
def _to_jpeg_data_url(image_bytes: bytes) -> str:
    """Normalise + downscale any image to a small JPEG data URL."""
    from PIL import Image, ImageOps
    img = Image.open(io.BytesIO(image_bytes))
    img = ImageOps.exif_transpose(img)          # fixes sideways phone screenshots
    if img.mode != "RGB":
        img = img.convert("RGB")
    if max(img.size) > MAX_IMAGE_SIDE:
        ratio = MAX_IMAGE_SIDE / max(img.size)
        img = img.resize((int(img.width * ratio), int(img.height * ratio)), Image.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85, optimize=True)
    return "data:image/jpeg;base64," + base64.b64encode(buf.getvalue()).decode()


def _clean_vision_output(text: str) -> str:
    """Strip reasoning tags / markdown fences some models emit."""
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"```[a-zA-Z]*|```", "", text)
    return text.strip()


def _vision_transcribe(data_url: str) -> str:
    """Send one image to Groq vision, trying each model until one answers."""
    last_error = None
    for model_id in VISION_MODELS:
        try:
            resp = client.chat.completions.create(
                model=model_id,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": OCR_PROMPT},
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                }],
                temperature=0,
                max_completion_tokens=1024,
                timeout=45,
            )
            text = _clean_vision_output(resp.choices[0].message.content or "")
            if "NO_TEXT_FOUND" in text.upper():
                return ""
            print(f"[OCR] transcribed via {model_id} ({len(text)} chars)")
            return text
        except Exception as e:
            last_error = e
            print(f"[OCR] model {model_id} failed: {e}")
            continue
    raise HTTPException(
        status_code=503,
        detail=f"Image reading is temporarily unavailable. Please paste the text instead. ({last_error})"
    )


def extract_text_from_image_bytes(image_bytes: bytes) -> str:
    if not image_bytes:
        raise HTTPException(status_code=422, detail="The uploaded image is empty.")
    try:
        data_url = _to_jpeg_data_url(image_bytes)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"That file isn't a readable image: {e}")
    return _vision_transcribe(data_url)


def _pdf_pages_to_jpegs(pdf_bytes: bytes, max_pages: int = 3) -> list[str]:
    """Render PDF pages to JPEG data URLs using pypdfium2 (no poppler needed)."""
    import pypdfium2 as pdfium
    pdf = pdfium.PdfDocument(pdf_bytes)
    urls = []
    for i in range(min(len(pdf), max_pages)):
        pil = pdf[i].render(scale=2).to_pil()
        buf = io.BytesIO()
        pil.convert("RGB").save(buf, format="JPEG", quality=85, optimize=True)
        urls.append(_to_jpeg_data_url(buf.getvalue()))
    return urls

# ── core analysis ─────────────────────────────────────────────────────────────
async def run_analysis(message: str) -> dict:
    user_content = f"""{FEW_SHOT_EXAMPLES}

Now analyze this message:
Message: "{message}"

Return exactly this JSON structure:
{{
  "risk_score": <integer 0-100>,
  "verdict": "<Safe | Suspicious | Likely Scam | Definite Scam>",
  "scam_type": "<null or one of: Registration Fee Scam, Guaranteed Job Scam, Work From Home Scam, Data Harvesting Scam, Impersonation Scam, Other>",
  "red_flags": ["<flag1>", "<flag2>"],
  "safe_signals": ["<signal1>"],
  "advice": "<one actionable sentence for the user>"
}}"""
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_content},
            ],
            temperature=0.1, timeout=10,
        )
        text = response.choices[0].message.content.strip()
        text = re.sub(r"```json|```", "", text).strip()
        return json.loads(text)
    except Exception as e:
        print(f"[FALLBACK TRIGGERED] Reason: {e}")
        return fallback_response(message)

MIN_MESSAGE_LENGTH = 10

# ── File type helpers ─────────────────────────────────────────────────────────
IMAGE_EXTS = {'jpg', 'jpeg', 'png', 'webp', 'bmp', 'tiff', 'tif'}
DOC_EXTS = {
    'pdf': 'application/pdf',
    'txt': 'text/plain',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'doc': 'application/msword'
}

# ── endpoint 1: text analysis ─────────────────────────────────────────────────
@app.post("/analyze")
@limiter.limit("10/minute")
async def analyze_message(request: Request, req: MessageRequest):
    message = (req.message or "").strip()
    if not message:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Message cannot be empty.")
    if len(message) < MIN_MESSAGE_LENGTH:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"Message is too short. Please provide at least {MIN_MESSAGE_LENGTH} characters.")
    return await run_analysis(message)

# ── endpoint 2: file upload (PDF + DOCX + TXT + Images) ──────────────────────
@app.post("/analyze-file")
@limiter.limit("5/minute")
async def analyze_file(request: Request, file: UploadFile = File(...)):
    content_type = file.content_type or ""
    filename = file.filename or ""
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    # Detect content type by extension if browser sends octet-stream
    if not content_type or content_type == 'application/octet-stream':
        if ext in IMAGE_EXTS:
            content_type = f'image/{ext}'
        elif ext in DOC_EXTS:
            content_type = DOC_EXTS[ext]

    allowed = ["application/pdf", "text/plain",
               "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
               "application/msword", "image/"]

    if not any(a in content_type for a in allowed) and ext not in IMAGE_EXTS and ext not in DOC_EXTS:
        raise HTTPException(status_code=400, detail="Unsupported file type. Upload PDF, DOCX, TXT, or image (JPG/PNG/WEBP).")

    contents = await file.read()
    if len(contents) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large. Please upload a file under 8MB.")
    extracted_text = ""

    # ── PDF ──
    if "pdf" in content_type or ext == 'pdf':
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(contents)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"

            # A PDF can carry a thin text layer (a heading, a watermark) while the
            # real content is scanned screenshots. Empty-check alone misses those,
            # so OCR whenever the text density per page is too low to be real.
            MIN_CHARS_PER_PAGE = 120
            page_count = max(len(pdfplumber.open(io.BytesIO(contents)).pages), 1)
            if len(extracted_text.strip()) < MIN_CHARS_PER_PAGE * page_count:
                try:
                    ocr_text = ""
                    for data_url in _pdf_pages_to_jpegs(contents, max_pages=5):
                        page_text = _vision_transcribe(data_url)
                        if page_text:
                            ocr_text += page_text + "\n"
                    if ocr_text.strip():
                        extracted_text = (extracted_text + "\n" + ocr_text).strip()
                except HTTPException:
                    raise
                except Exception as ocr_e:
                    if not extracted_text.strip():
                        raise HTTPException(status_code=422, detail=f"This PDF appears image-based and could not be read: {str(ocr_e)}")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Could not process PDF: {str(e)}")

    # ── DOCX ──
    elif "wordprocessingml" in content_type or "msword" in content_type or ext in ('docx', 'doc'):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(contents))
            extracted_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Could not read Word document: {str(e)}")

    # ── Plain text ──
    elif "text/plain" in content_type or ext == 'txt':
        try:
            extracted_text = contents.decode("utf-8")
        except UnicodeDecodeError:
            extracted_text = contents.decode("latin-1")

    # ── Images ──
    elif "image/" in content_type or ext in IMAGE_EXTS:
        extracted_text = extract_text_from_image_bytes(contents)
        if not extracted_text:
            raise HTTPException(status_code=422, detail="No text found in image. Make sure the image contains readable text.")

    extracted_text = extracted_text.strip()
    if not extracted_text:
        raise HTTPException(status_code=422, detail="No text could be extracted. The file may be empty or unreadable.")
    if len(extracted_text) < MIN_MESSAGE_LENGTH:
        raise HTTPException(status_code=400, detail="Extracted text is too short to analyze.")
    if len(extracted_text) > 3000:
        extracted_text = extracted_text[:3000] + "..."

    result = await run_analysis(extracted_text)
    result["extracted_preview"] = extracted_text[:300] + ("..." if len(extracted_text) > 300 else "")
    result["source"] = f"Extracted from: {file.filename}"
    return result

# ── endpoint 3: report a scam ─────────────────────────────────────────────────
@app.post("/report-scam")
@limiter.limit("3/minute")
async def report_scam(request: Request, req: ReportRequest):
    if not SUPABASE_URL or not SUPABASE_KEY:
        raise HTTPException(status_code=500, detail="Database not configured.")
    try:
        async with httpx.AsyncClient() as client_http:
            response = await client_http.post(
                f"{SUPABASE_URL}/rest/v1/scam_reports",
                headers={"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}", "Content-Type": "application/json", "Prefer": "return=minimal"},
                json={"message_preview": req.message_preview[:200], "scam_type": req.scam_type, "risk_score": req.risk_score, "verdict": req.verdict, "red_flags": req.red_flags}
            )
        if response.status_code in [200, 201]:
            return {"success": True, "message": "Scam reported successfully. Thank you for helping others!"}
        else:
            raise HTTPException(status_code=500, detail="Failed to save report.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

# ── endpoint 4: get recent scam reports ───────────────────────────────────────
@app.get("/recent-scams")
@limiter.limit("30/minute")
async def get_recent_scams(request: Request):
    if not SUPABASE_URL or not SUPABASE_KEY:
        raise HTTPException(status_code=500, detail="Database not configured.")
    try:
        async with httpx.AsyncClient() as client_http:
            response = await client_http.get(
                f"{SUPABASE_URL}/rest/v1/scam_reports",
                headers={"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}"},
                params={"select": "id,message_preview,scam_type,risk_score,verdict,red_flags,reported_at", "order": "reported_at.desc", "limit": "20"}
            )
        return response.json()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

# ── endpoint 5: chat ──────────────────────────────────────────────────────────
@app.post("/chat")
@limiter.limit("20/minute")
async def chat(request: Request, req: ChatRequest):
    if not req.session_id or not req.message.strip():
        raise HTTPException(status_code=400, detail="session_id and message are required.")

    history = []
    if SUPABASE_URL and SUPABASE_KEY:
        try:
            async with httpx.AsyncClient() as client_http:
                resp = await client_http.get(
                    f"{SUPABASE_URL}/rest/v1/chat_history",
                    headers={"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}"},
                    params={"select": "role,message,created_at", "session_id": f"eq.{req.session_id}", "order": "created_at.asc", "limit": "10"}
                )
            history = resp.json() if resp.status_code == 200 else []
        except:
            history = []

    messages = [{"role": "system", "content": CHAT_SYSTEM_PROMPT}]
    for h in history:
        messages.append({"role": h["role"], "content": h["message"]})
    messages.append({"role": "user", "content": req.message})

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=0.3,
            timeout=15,
            max_tokens=800,
        )
        reply = response.choices[0].message.content.strip()
    except Exception as e:
        reply = "I'm having trouble connecting right now. Please try again in a moment."

    if SUPABASE_URL and SUPABASE_KEY:
        try:
            async with httpx.AsyncClient() as client_http:
                await client_http.post(
                    f"{SUPABASE_URL}/rest/v1/chat_history",
                    headers={"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}", "Content-Type": "application/json", "Prefer": "return=minimal"},
                    json={"session_id": req.session_id, "role": "user", "message": req.message}
                )
                await client_http.post(
                    f"{SUPABASE_URL}/rest/v1/chat_history",
                    headers={"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}", "Content-Type": "application/json", "Prefer": "return=minimal"},
                    json={"session_id": req.session_id, "role": "assistant", "message": reply}
                )
        except:
            pass

    return {"reply": reply, "session_id": req.session_id}

# ── endpoint 6: get chat history ──────────────────────────────────────────────
@app.get("/chat-history/{session_id}")
@limiter.limit("30/minute")
async def get_chat_history(request: Request, session_id: str):
    if not SUPABASE_URL or not SUPABASE_KEY:
        return []
    try:
        async with httpx.AsyncClient() as client_http:
            resp = await client_http.get(
                f"{SUPABASE_URL}/rest/v1/chat_history",
                headers={"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}"},
                params={"select": "role,message,created_at", "session_id": f"eq.{session_id}", "order": "created_at.asc", "limit": "50"}
            )
        return resp.json() if resp.status_code == 200 else []
    except:
        return []

# ── health check ─────────────────────────────────────────────────────────────
@app.get("/")
async def root():
    return {
        "status": "ScamShield API is running",
        "version": "3.2",
        "ocr": "groq-vision",
        "ocr_model": VISION_MODELS[0],
    }